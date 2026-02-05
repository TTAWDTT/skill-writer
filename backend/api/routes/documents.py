"""
Documents API 路由
处理生成的文档
"""
from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel
from typing import Optional, List
from datetime import datetime
import uuid
import io
import re
import base64
import tempfile
import httpx

from backend.models.database import get_database, Document as DocumentModel

router = APIRouter()
db = get_database()



class DocumentCreate(BaseModel):
    """创建文档请求"""
    title: str
    skill_id: str
    content: str
    session_id: Optional[str] = None


class DocumentUpdate(BaseModel):
    """更新文档请求"""
    title: Optional[str] = None
    content: Optional[str] = None


class DocumentResponse(BaseModel):
    """文档响应"""
    id: str
    title: str
    skill_id: str
    content: str
    created_at: str
    updated_at: str


class ExportRequest(BaseModel):
    """导出请求"""
    content: str
    format: str  # md, docx, pdf
    filename: str


@router.post("/", response_model=DocumentResponse)
async def create_document(doc: DocumentCreate):
    """创建新文档"""
    doc_id = str(uuid.uuid4())
    now = datetime.utcnow()

    with db.get_session() as db_session:
        record = DocumentModel(
            id=doc_id,
            session_id=doc.session_id,
            skill_id=doc.skill_id,
            title=doc.title,
            content=doc.content,
            created_at=now,
            updated_at=now,
        )
        db_session.add(record)
        db_session.commit()

    return DocumentResponse(
        id=doc_id,
        title=doc.title,
        skill_id=doc.skill_id,
        content=doc.content,
        created_at=now.isoformat(),
        updated_at=now.isoformat(),
    )


@router.get("/", response_model=List[DocumentResponse])
async def list_documents(skill_id: Optional[str] = None):
    """获取文档列表"""
    with db.get_session() as db_session:
        query = db_session.query(DocumentModel)
        if skill_id:
            query = query.filter(DocumentModel.skill_id == skill_id)
        records = query.order_by(DocumentModel.updated_at.desc()).all()

    return [
        DocumentResponse(
            id=doc.id,
            title=doc.title,
            skill_id=doc.skill_id,
            content=doc.content,
            created_at=doc.created_at.isoformat() if doc.created_at else "",
            updated_at=doc.updated_at.isoformat() if doc.updated_at else "",
        )
        for doc in records
    ]


@router.get("/{doc_id}", response_model=DocumentResponse)
async def get_document(doc_id: str):
    """获取单个文档"""
    with db.get_session() as db_session:
        record = db_session.query(DocumentModel).filter(DocumentModel.id == doc_id).first()
        if not record:
            raise HTTPException(status_code=404, detail=f"Document not found: {doc_id}")

    return DocumentResponse(
        id=record.id,
        title=record.title,
        skill_id=record.skill_id,
        content=record.content,
        created_at=record.created_at.isoformat() if record.created_at else "",
        updated_at=record.updated_at.isoformat() if record.updated_at else "",
    )


@router.put("/{doc_id}", response_model=DocumentResponse)
async def update_document(doc_id: str, update: DocumentUpdate):
    """更新文档"""
    with db.get_session() as db_session:
        record = db_session.query(DocumentModel).filter(DocumentModel.id == doc_id).first()
        if not record:
            raise HTTPException(status_code=404, detail=f"Document not found: {doc_id}")

        if update.title is not None:
            record.title = update.title
        if update.content is not None:
            record.content = update.content

        record.updated_at = datetime.utcnow()
        db_session.commit()

        return DocumentResponse(
            id=record.id,
            title=record.title,
            skill_id=record.skill_id,
            content=record.content,
            created_at=record.created_at.isoformat() if record.created_at else "",
            updated_at=record.updated_at.isoformat() if record.updated_at else "",
        )


@router.delete("/{doc_id}")
async def delete_document(doc_id: str):
    """删除文档"""
    with db.get_session() as db_session:
        record = db_session.query(DocumentModel).filter(DocumentModel.id == doc_id).first()
        if not record:
            raise HTTPException(status_code=404, detail=f"Document not found: {doc_id}")
        db_session.delete(record)
        db_session.commit()

    return {"message": "Document deleted"}


@router.get("/{doc_id}/export")
async def export_document(doc_id: str, format: str = "markdown"):
    """导出文档"""
    with db.get_session() as db_session:
        record = db_session.query(DocumentModel).filter(DocumentModel.id == doc_id).first()
        if not record:
            raise HTTPException(status_code=404, detail=f"Document not found: {doc_id}")

    doc = {
        "title": record.title,
        "content": record.content,
    }

    if format == "markdown":
        return {
            "format": "markdown",
            "content": doc["content"],
            "filename": f"{doc['title']}.md",
        }
    elif format == "html":
        # 简单的 Markdown 转 HTML（生产环境应使用 markdown 库）
        html_content = f"<html><body><h1>{doc['title']}</h1><pre>{doc['content']}</pre></body></html>"
        return {
            "format": "html",
            "content": html_content,
            "filename": f"{doc['title']}.html",
        }
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {format}")


@router.post("/export")
async def export_content(request: ExportRequest):
    """导出内容为指定格式"""
    content = request.content
    format_type = request.format.lower()
    filename = request.filename

    if format_type == "md":
        return Response(
            content=content.encode('utf-8'),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{filename}.md"'}
        )

    elif format_type == "docx":
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from docx.enum.style import WD_STYLE_TYPE
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # 设置文档默认字体
            def set_document_default_font(doc, font_name='宋体'):
                """设置文档级别的默认字体"""
                # 设置 Normal 样式的字体
                style = doc.styles['Normal']
                style.font.name = font_name

                # 确保 rPr 存在
                rPr = style._element.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.insert(0, rFonts)
                rFonts.set(qn('w:eastAsia'), font_name)
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)

                # 设置文档的默认字体
                doc_defaults = doc.styles.element
                rPrDefault = doc_defaults.find(qn('w:docDefaults'))
                if rPrDefault is None:
                    rPrDefault = OxmlElement('w:docDefaults')
                    doc_defaults.insert(0, rPrDefault)

                rPrDefaultRPr = rPrDefault.find(qn('w:rPrDefault'))
                if rPrDefaultRPr is None:
                    rPrDefaultRPr = OxmlElement('w:rPrDefault')
                    rPrDefault.append(rPrDefaultRPr)

                rPr = rPrDefaultRPr.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    rPrDefaultRPr.append(rPr)

                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)

                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rFonts.set(qn('w:eastAsia'), font_name)
                rFonts.set(qn('w:cs'), font_name)

            # 应用文档默认字体
            set_document_default_font(doc, '宋体')

            # 设置默认字体为宋体
            def set_run_font(run, font_name='宋体', font_size=12, bold=False, italic=False):
                """设置 run 的字体"""
                run.font.size = Pt(font_size)
                run.font.bold = bold
                run.font.italic = italic

                # 设置字体名称（西文）
                run.font.name = font_name

                # 确保 rPr 元素存在
                rPr = run._element.get_or_add_rPr()

                # 创建或获取 rFonts 元素
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.insert(0, rFonts)

                # 设置所有字体属性
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rFonts.set(qn('w:eastAsia'), font_name)
                rFonts.set(qn('w:cs'), font_name)

            def parse_inline_formatting(paragraph, text, font_size=12):
                """解析并应用行内格式（加粗、斜体）"""
                # 正则表达式匹配加粗和斜体
                # **text** 或 __text__ 为加粗
                # *text* 或 _text_ 为斜体
                # ***text*** 为加粗斜体
                pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|__(.+?)__|(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|(?<!_)_(?!_)(.+?)(?<!_)_(?!_))'

                last_end = 0
                for match in re.finditer(pattern, text):
                    # 添加匹配前的普通文本
                    if match.start() > last_end:
                        run = paragraph.add_run(text[last_end:match.start()])
                        set_run_font(run, font_size=font_size)

                    # 判断是哪种格式
                    if match.group(2):  # ***text*** - 加粗斜体
                        run = paragraph.add_run(match.group(2))
                        set_run_font(run, font_size=font_size, bold=True, italic=True)
                    elif match.group(3):  # **text** - 加粗
                        run = paragraph.add_run(match.group(3))
                        set_run_font(run, font_size=font_size, bold=True)
                    elif match.group(4):  # __text__ - 加粗
                        run = paragraph.add_run(match.group(4))
                        set_run_font(run, font_size=font_size, bold=True)
                    elif match.group(5):  # *text* - 斜体
                        run = paragraph.add_run(match.group(5))
                        set_run_font(run, font_size=font_size, italic=True)
                    elif match.group(6):  # _text_ - 斜体
                        run = paragraph.add_run(match.group(6))
                        set_run_font(run, font_size=font_size, italic=True)

                    last_end = match.end()

                # 添加剩余的普通文本
                if last_end < len(text):
                    run = paragraph.add_run(text[last_end:])
                    set_run_font(run, font_size=font_size)

                # 如果没有任何匹配，添加整个文本
                if last_end == 0 and text:
                    run = paragraph.add_run(text)
                    set_run_font(run, font_size=font_size)

            image_pattern = re.compile(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$')

            def _load_image_bytes(src: str) -> Optional[tuple[bytes, str]]:
                if not src:
                    return None
                src = src.strip()

                # data URI
                if src.startswith("data:image/"):
                    m = re.match(r"^data:(image/[a-zA-Z0-9.+-]+);base64,(.+)$", src)
                    if not m:
                        return None
                    mime = m.group(1).lower()
                    b64 = m.group(2)
                    try:
                        raw = base64.b64decode(b64)
                    except Exception:
                        return None
                    ext = "png"
                    if "jpeg" in mime or "jpg" in mime:
                        ext = "jpg"
                    elif "png" in mime:
                        ext = "png"
                    return raw, ext

                # remote
                if src.startswith("http://") or src.startswith("https://"):
                    try:
                        with httpx.Client(timeout=30.0, follow_redirects=True) as client:
                            resp = client.get(src)
                            resp.raise_for_status()
                            raw = resp.content
                        # best-effort ext guess
                        ext = "png" if src.lower().endswith(".png") else "jpg" if src.lower().endswith((".jpg", ".jpeg")) else "png"
                        return raw, ext
                    except Exception:
                        return None

                return None

            def _add_markdown_image(line: str) -> bool:
                m = image_pattern.match(line.strip())
                if not m:
                    return False
                alt = (m.group(1) or "").strip()
                src = (m.group(2) or "").strip()
                loaded = _load_image_bytes(src)
                if not loaded:
                    # fall back: treat as plain text
                    para = doc.add_paragraph()
                    parse_inline_formatting(para, line)
                    return True

                raw, ext = loaded
                stream = io.BytesIO(raw)
                # Help python-docx with type inference
                try:
                    stream.name = f"image.{ext}"  # type: ignore[attr-defined]
                except Exception:
                    pass

                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                # Fit page width roughly (A4 default)
                run.add_picture(stream, width=Inches(6.2))
                if alt:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = cap.add_run(alt)
                    set_run_font(r, font_name="宋体", font_size=10, bold=False)
                return True

            # 解析 Markdown 并转换为 Word 格式
            lines = content.split('\n')
            i = 0
            while i < len(lines):
                line = lines[i].strip()

                if not line:
                    i += 1
                    continue

                # Markdown image: ![alt](src)
                if _add_markdown_image(line):
                    i += 1
                    continue

                # 处理标题 - 使用普通段落而不是heading样式，以便更好控制字体
                if line.startswith('#'):
                    level = len(re.match(r'^#+', line).group())
                    title_text = line.lstrip('#').strip()
                    # 使用普通段落，手动设置格式
                    heading_para = doc.add_paragraph()
                    # 设置标题字体大小
                    sizes = {1: 26, 2: 20, 3: 18, 4: 16, 5: 14, 6: 12}
                    run = heading_para.add_run(title_text)
                    set_run_font(run, font_name='黑体', font_size=sizes.get(level, 14), bold=True)
                    # 设置段落间距
                    heading_para.paragraph_format.space_before = Pt(12)
                    heading_para.paragraph_format.space_after = Pt(6)
                    i += 1

                # 处理无序列表
                elif line.startswith('- ') or (line.startswith('* ') and not line.startswith('**')):
                    list_text = line[2:]
                    para = doc.add_paragraph(style='List Bullet')
                    parse_inline_formatting(para, list_text)
                    i += 1

                # 处理有序列表
                elif re.match(r'^\d+\. ', line):
                    list_text = re.sub(r'^\d+\. ', '', line)
                    para = doc.add_paragraph(style='List Number')
                    parse_inline_formatting(para, list_text)
                    i += 1

                # 普通段落
                else:
                    # 收集连续的非空行作为一个段落
                    para_lines = [line]
                    i += 1
                    while (
                        i < len(lines)
                        and lines[i].strip()
                        and not lines[i].strip().startswith('#')
                        and not image_pattern.match(lines[i].strip())
                        and not lines[i].strip().startswith('- ')
                        and not (lines[i].strip().startswith('* ') and not lines[i].strip().startswith('**'))
                        and not re.match(r'^\d+\. ', lines[i].strip())
                    ):
                        para_lines.append(lines[i].strip())
                        i += 1

                    para = doc.add_paragraph()
                    full_text = ' '.join(para_lines)
                    parse_inline_formatting(para, full_text)

            # 保存到内存
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return Response(
                content=buffer.getvalue(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{filename}.docx"'}
            )

        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="python-docx not installed. Run: pip install python-docx"
            )

    elif format_type == "pdf":
        try:
            from fpdf import FPDF
            import re as re_module
            import os

            # 创建 PDF，设置合理的边距
            pdf = FPDF()
            pdf.set_margins(15, 15, 15)  # 左、上、右边距
            pdf.set_auto_page_break(auto=True, margin=15)  # 自动分页，底部边距
            pdf.add_page()

            # 添加中文字体支持
            # 尝试使用系统字体
            font_loaded = False
            font_name = 'Helvetica'  # FPDF 核心字体
            try:
                pdf.add_font('SimSun', '', 'C:/Windows/Fonts/simsun.ttc', uni=True)
                font_name = 'SimSun'
                font_loaded = True
            except Exception as e1:
                try:
                    pdf.add_font('Microsoft YaHei', '', 'C:/Windows/Fonts/msyh.ttc', uni=True)
                    font_name = 'Microsoft YaHei'
                    font_loaded = True
                except Exception as e2:
                    # 尝试其他常见中文字体路径
                    try:
                        pdf.add_font('SimHei', '', 'C:/Windows/Fonts/simhei.ttf', uni=True)
                        font_name = 'SimHei'
                        font_loaded = True
                    except Exception:
                        # 使用 Helvetica 作为最后的后备方案
                        font_name = 'Helvetica'
                        font_loaded = False

            pdf.set_font(font_name, size=12)

            # 清理内容中可能导致问题的字符
            def clean_text(text):
                """清理可能导致PDF渲染问题的字符"""
                # 替换一些特殊字符
                text = text.replace('\r\n', '\n').replace('\r', '\n')
                # 移除可能导致问题的控制字符
                text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t')
                return text

            content = clean_text(content)

            image_pattern = re_module.compile(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$')

            def _load_image_bytes(src: str) -> Optional[tuple[bytes, str]]:
                if not src:
                    return None
                src = src.strip()
                if src.startswith("data:image/"):
                    m = re_module.match(r"^data:(image/[a-zA-Z0-9.+-]+);base64,(.+)$", src)
                    if not m:
                        return None
                    mime = m.group(1).lower()
                    b64 = m.group(2)
                    try:
                        raw = base64.b64decode(b64)
                    except Exception:
                        return None
                    ext = "png"
                    if "jpeg" in mime or "jpg" in mime:
                        ext = "jpg"
                    elif "png" in mime:
                        ext = "png"
                    return raw, ext

                if src.startswith("http://") or src.startswith("https://"):
                    try:
                        with httpx.Client(timeout=30.0, follow_redirects=True) as client:
                            resp = client.get(src)
                            resp.raise_for_status()
                            raw = resp.content
                        ext = "png" if src.lower().endswith(".png") else "jpg" if src.lower().endswith((".jpg", ".jpeg")) else "png"
                        return raw, ext
                    except Exception:
                        return None
                return None

            # 解析 Markdown 并添加到 PDF
            with tempfile.TemporaryDirectory(prefix="skillwriter_pdf_") as td:
                lines = content.split('\n')
                for line in lines:
                    line = line.rstrip()

                    if not line:
                        pdf.ln(5)
                        continue

                    # Image line
                    im = image_pattern.match(line.strip())
                    if im:
                        alt = (im.group(1) or "").strip()
                        src = (im.group(2) or "").strip()
                        loaded = _load_image_bytes(src)
                        if loaded:
                            raw, ext = loaded
                            tmp_path = os.path.join(td, f"img_{uuid.uuid4().hex}.{ext}")
                            with open(tmp_path, "wb") as f:
                                f.write(raw)
                            # Page break if needed
                            if pdf.get_y() > pdf.page_break_trigger - 60:
                                pdf.add_page()
                                pdf.set_font(font_name, size=12)
                            pdf.ln(3)
                            max_w = pdf.w - pdf.l_margin - pdf.r_margin
                            try:
                                pdf.image(tmp_path, w=max_w)
                            except Exception:
                                pass
                            if alt:
                                pdf.ln(2)
                                pdf.set_font_size(10)
                                pdf.multi_cell(0, 5, alt)
                                pdf.set_font_size(12)
                            pdf.ln(4)
                            continue

                    # 如果没有加载中文字体，跳过非ASCII字符的行（避免渲染错误）
                    if not font_loaded:
                        try:
                            line.encode('latin-1')
                        except UnicodeEncodeError:
                            line = ''.join(char for char in line if ord(char) < 128)
                            if not line.strip():
                                pdf.ln(3)
                                continue

                    try:
                        # 处理标题
                        if line.startswith('#'):
                            level = len(re_module.match(r'^#+', line).group())
                            title_text = line.lstrip('#').strip()
                            if not title_text:
                                continue
                            sizes = {1: 18, 2: 16, 3: 14, 4: 12, 5: 11, 6: 10}
                            pdf.set_font_size(sizes.get(level, 12))
                            pdf.ln(6)
                            pdf.multi_cell(0, 8, title_text)
                            pdf.set_font_size(12)
                            pdf.ln(3)

                        # 处理列表
                        elif line.startswith('- ') or line.startswith('* '):
                            list_text = line[2:].strip()
                            if list_text:
                                pdf.multi_cell(0, 6, '  - ' + list_text)

                        elif re_module.match(r'^\d+\. ', line):
                            pdf.multi_cell(0, 6, '  ' + line)

                        # 普通段落
                        else:
                            pdf.multi_cell(0, 6, line)

                    except Exception:
                        try:
                            simple_line = ''.join(char for char in line if ord(char) < 128)
                            if simple_line.strip():
                                pdf.multi_cell(0, 6, simple_line)
                        except Exception:
                            pass

            # 输出 PDF
            pdf_content = pdf.output()

            return Response(
                content=bytes(pdf_content),
                media_type="application/pdf",
                headers={"Content-Disposition": f'attachment; filename="{filename}.pdf"'}
            )

        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="PDF export requires fpdf2. Run: pip install fpdf2"
            )
        except Exception as e:
            import traceback
            error_detail = f"PDF generation failed: {str(e)}\n{traceback.format_exc()}"
            print(f"[PDF Export Error] {error_detail}")
            raise HTTPException(
                status_code=500,
                detail=f"PDF generation failed: {str(e)}"
            )

    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {format_type}")
